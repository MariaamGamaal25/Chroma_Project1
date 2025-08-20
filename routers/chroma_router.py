# import os
# import shutil
# import tempfile
# from fastapi import APIRouter, HTTPException, UploadFile, File
# from typing import Any
# from dtos.document import Document
# from utils.enums import Endpoints, Messages, ChromaDBConfig, MetadataKeys
# from chroma_functionalities import get_chroma_collection, sync_text_file_with_collection
# from utils.logger import setup_logger
# from PyPDF2 import PdfReader
# from docx import Document as DocxDocument

# logger = setup_logger(__name__)
# router = APIRouter()

# @router.get(Endpoints.DATA, response_model=list[Document], tags=["Data Management"])
# def get_all_data():
#     """Retrieve all documents from the ChromaDB collection."""
#     try:
#         collection = get_chroma_collection()
#         results = collection.get(ids=None, include=['documents', 'metadatas'])
#         return [
#             Document(
#                 id=results['ids'][i],
#                 text=results['documents'][i],
#                 metadata=results['metadatas'][i]
#             )
#             for i in range(len(results.get('ids', [])))
#         ]
#     except Exception as e:
#         logger.error(Messages.ERROR_GENERAL.format(error=str(e)))
#         raise HTTPException(status_code=500, detail=str(e))


# @router.get(Endpoints.DATA_BY_ID, response_model=Document, tags=["Data Management"])
# def get_document_by_id(doc_id: str):
#     """Retrieve a single document by ID."""
#     try:
#         collection = get_chroma_collection()
#         results = collection.get(ids=[doc_id], include=['documents', 'metadatas'])
#         if not results.get('ids'):
#             raise HTTPException(status_code=404, detail=Messages.DOC_NOT_FOUND)
#         return Document(
#             id=results['ids'][0],
#             text=results['documents'][0],
#             metadata=results['metadatas'][0]
#         )
#     except HTTPException:
#         raise
#     except Exception as e:
#         logger.error(Messages.ERROR_GENERAL.format(error=str(e)))
#         raise HTTPException(status_code=500, detail=str(e))


# @router.get(Endpoints.RAW_DATA, tags=["Data Management"])
# def get_raw_data() -> dict[str, Any]:
#     """Retrieve all documents from ChromaDB in raw format."""
#     try:
#         collection = get_chroma_collection()
#         results = collection.get(ids=None, include=['documents', 'metadatas'])
#         return {
#             'documents': results.get('documents', []),
#             'ids': results.get('ids', []),
#             'metadatas': results.get('metadatas', []),
#             'embeddings': None,
#             'distances': None,
#             'uris': None,
#             'data': None,
#         }
#     except Exception as e:
#         logger.error(Messages.ERROR_GENERAL.format(error=str(e)))
#         raise HTTPException(status_code=500, detail=str(e))


# @router.delete(Endpoints.DATA_BY_ID, tags=["Data Management"])
# def delete_document_by_id(doc_id: str):
#     """Delete a document by ID."""
#     try:
#         collection = get_chroma_collection()
#         results = collection.get(ids=[doc_id])
#         if not results['ids']:
#             raise HTTPException(status_code=404, detail=Messages.DOC_NOT_FOUND)
#         collection.delete(ids=[doc_id])
#         return {"message": Messages.DOC_DELETED.format(doc_id=doc_id)}
#     except HTTPException:
#         raise
#     except Exception as e:
#         logger.error(Messages.ERROR_DELETION.format(error=str(e)))
#         raise HTTPException(status_code=500, detail=str(e))


# @router.delete(Endpoints.DATA, tags=["Data Management"])
# def delete_all_data():
#     """Delete all documents from ChromaDB."""
#     try:
#         collection = get_chroma_collection()
#         count_before_delete = collection.count()
#         if count_before_delete == 0:
#             return {"message": "No documents to delete"}
#         ids = collection.get(ids=None, include=[])['ids']
#         if ids:
#             collection.delete(ids=ids)
#         count_after_delete = collection.count()
#         if count_after_delete == 0:
#             return {"message": Messages.ALL_DOCS_DELETED.format(count_before_delete=count_before_delete)}
#         else:
#             raise HTTPException(status_code=500, detail=Messages.DELETE_FAILED)
#     except Exception as e:
#         logger.error(Messages.ERROR_DELETION.format(error=str(e)))
#         raise HTTPException(status_code=500, detail=str(e))


# @router.post(Endpoints.UPLOAD_TEXT_FILE, tags=["Sync Data to chroma"])
# async def upload_text_file(file: UploadFile = File(...), force: bool = False):
#     """
#     Upload a text, PDF, or DOCX file and store its contents in ChromaDB.
#     """
#     try:
#         # Validate extension
#         file_extension = os.path.splitext(file.filename)[1].lower()
#         if file_extension not in ('.txt', '.pdf', '.docx'):
#             raise HTTPException(status_code=400, detail=Messages.INVALID_FILE_TYPE)

#         # Create a temporary directory to handle the file
#         with tempfile.TemporaryDirectory() as temp_dir:
#             temp_file_path = os.path.join(temp_dir, file.filename)

#             # Save the uploaded file to the temporary directory
#             with open(temp_file_path, "wb") as temp_file:
#                 shutil.copyfileobj(file.file, temp_file)

#             # Extract text depending on file type
#             extracted_text = ""
#             if file_extension == ".txt":
#                 with open(temp_file_path, "r", encoding="utf-8") as f:
#                     extracted_text = f.read()

#             elif file_extension == ".pdf":
#                 reader = PdfReader(temp_file_path)
#                 extracted_text = "\n".join([page.extract_text() or "" for page in reader.pages])

#             elif file_extension == ".docx":
#                 try:
#                     doc = DocxDocument(temp_file_path)
#                     extracted_text = "\n".join([p.text for p in doc.paragraphs])
#                 except Exception as docx_error:
#                     logger.error(f"Failed to process DOCX file: {docx_error}")
#                     raise HTTPException(status_code=500, detail=f"Failed to process DOCX file: {str(docx_error)}")

#             # Pass the extracted text to the sync function
#             collection = get_chroma_collection()
#             sync_text_file_with_collection(collection, temp_file_path, force=force)
            
#             return {"message": Messages.FILE_UPLOADED.format(file_name=file.filename)}

#     except FileNotFoundError as e:
#         raise HTTPException(status_code=400, detail=str(e))
#     except Exception as e:
#         logger.error(Messages.ERROR_FILE_UPLOAD.format(error=str(e)))
#         raise HTTPException(status_code=500, detail=str(e))

# @router.delete(Endpoints.DROP_DATABASE, tags=["Data Management"])
# def drop_database():
#     """Delete the entire ChromaDB database folder."""
#     try:
#         db_directory = os.path.abspath(ChromaDBConfig.DB_DIRECTORY)
#         if not os.path.exists(db_directory):
#             raise HTTPException(
#                 status_code=404,
#                 detail=Messages.DATABASE_NOT_FOUND.format(db_directory=db_directory)
#             )
#         shutil.rmtree(db_directory, ignore_errors=False)
#         logger.info(Messages.DATABASE_DROPPED.format(db_directory=db_directory))
#         return {"message": Messages.DATABASE_DROPPED.format(db_directory=db_directory)}
#     except PermissionError as e:
#         logger.error(f"Permission denied while deleting DB folder: {str(e)}")
#         raise HTTPException(status_code=403, detail="Permission denied while deleting DB folder.")
#     except Exception as e:
#         logger.error(Messages.ERROR_GENERAL.format(error=str(e)))
#         raise HTTPException(status_code=500, detail=str(e))
import os
import shutil
import tempfile
from fastapi import APIRouter, HTTPException, UploadFile, File
from typing import Any
from dtos.document import Document
from utils.enums import Endpoints, Messages, ChromaDBConfig, MetadataKeys
from chroma_functionalities import get_chroma_collection, sync_text_file_with_collection, read_and_chunk_file
from utils.logger import setup_logger
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

logger = setup_logger(__name__)
router = APIRouter()

@router.get(Endpoints.DATA, response_model=list[Document], tags=["Data Management"])
def get_all_data():
    """Retrieve all documents from the ChromaDB collection."""
    try:
        collection = get_chroma_collection()
        results = collection.get(ids=None, include=['documents', 'metadatas'])
        return [
            Document(
                id=results['ids'][i],
                text=results['documents'][i],
                metadata=results['metadatas'][i]
            )
            for i in range(len(results.get('ids', [])))
        ]
    except Exception as e:
        logger.error(Messages.ERROR_GENERAL.format(error=str(e)))
        raise HTTPException(status_code=500, detail=str(e))


@router.get(Endpoints.DATA_BY_ID, response_model=Document, tags=["Data Management"])
def get_document_by_id(doc_id: str):
    """Retrieve a single document by ID."""
    try:
        collection = get_chroma_collection()
        results = collection.get(ids=[doc_id], include=['documents', 'metadatas'])
        if not results.get('ids'):
            raise HTTPException(status_code=404, detail=Messages.DOC_NOT_FOUND)
        return Document(
            id=results['ids'][0],
            text=results['documents'][0],
            metadata=results['metadatas'][0]
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(Messages.ERROR_GENERAL.format(error=str(e)))
        raise HTTPException(status_code=500, detail=str(e))


@router.get(Endpoints.RAW_DATA, tags=["Data Management"])
def get_raw_data() -> dict[str, Any]:
    """Retrieve all documents from ChromaDB in raw format."""
    try:
        collection = get_chroma_collection()
        results = collection.get(ids=None, include=['documents', 'metadatas'])
        return {
            'documents': results.get('documents', []),
            'ids': results.get('ids', []),
            'metadatas': results.get('metadatas', []),
            'embeddings': None,
            'distances': None,
            'uris': None,
            'data': None,
        }
    except Exception as e:
        logger.error(Messages.ERROR_GENERAL.format(error=str(e)))
        raise HTTPException(status_code=500, detail=str(e))


@router.delete(Endpoints.DATA_BY_ID, tags=["Data Management"])
def delete_document_by_id(doc_id: str):
    """Delete a document by ID."""
    try:
        collection = get_chroma_collection()
        results = collection.get(ids=[doc_id])
        if not results['ids']:
            raise HTTPException(status_code=404, detail=Messages.DOC_NOT_FOUND)
        collection.delete(ids=[doc_id])
        return {"message": Messages.DOC_DELETED.format(doc_id=doc_id)}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(Messages.ERROR_DELETION.format(error=str(e)))
        raise HTTPException(status_code=500, detail=str(e))


@router.delete(Endpoints.DATA, tags=["Data Management"])
def delete_all_data():
    """Delete all documents from ChromaDB."""
    try:
        collection = get_chroma_collection()
        count_before_delete = collection.count()
        if count_before_delete == 0:
            return {"message": "No documents to delete"}
        ids = collection.get(ids=None, include=[])['ids']
        if ids:
            collection.delete(ids=ids)
        count_after_delete = collection.count()
        if count_after_delete == 0:
            return {"message": Messages.ALL_DOCS_DELETED.format(count_before_delete=count_before_delete)}
        else:
            raise HTTPException(status_code=500, detail=Messages.DELETE_FAILED)
    except Exception as e:
        logger.error(Messages.ERROR_DELETION.format(error=str(e)))
        raise HTTPException(status_code=500, detail=str(e))


@router.post(Endpoints.UPLOAD_TEXT_FILE, tags=["Sync Data to chroma"])
async def upload_text_file(file: UploadFile = File(...), force: bool = False):
    """
    Upload a text, PDF, or DOCX file and store its contents in ChromaDB.
    """
    try:
        # Validate extension
        file_extension = os.path.splitext(file.filename)[1].lower()
        if file_extension not in ('.txt', '.pdf', '.docx'):
            raise HTTPException(status_code=400, detail=Messages.INVALID_FILE_TYPE)

        # Create a temporary directory to handle the file
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_file_path = os.path.join(temp_dir, file.filename)

            # Save the uploaded file to the temporary directory
            with open(temp_file_path, "wb") as temp_file:
                shutil.copyfileobj(file.file, temp_file)

            # Pass the temp file path directly to the `sync_text_file_with_collection` function
            collection = get_chroma_collection()
            sync_text_file_with_collection(collection, temp_file_path, force=force)
            
            return {"message": Messages.FILE_UPLOADED.format(file_name=file.filename)}

    except FileNotFoundError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(Messages.ERROR_FILE_UPLOAD.format(error=str(e)))
        raise HTTPException(status_code=500, detail=str(e))


@router.delete(Endpoints.DROP_DATABASE, tags=["Data Management"])
def drop_database():
    """Delete the entire ChromaDB database folder."""
    try:
        db_directory = os.path.abspath(ChromaDBConfig.DB_DIRECTORY)
        if not os.path.exists(db_directory):
            raise HTTPException(
                status_code=404,
                detail=Messages.DATABASE_NOT_FOUND.format(db_directory=db_directory)
            )
        shutil.rmtree(db_directory, ignore_errors=False)
        logger.info(Messages.DATABASE_DROPPED.format(db_directory=db_directory))
        return {"message": Messages.DATABASE_DROPPED.format(db_directory=db_directory)}
    except PermissionError as e:
        logger.error(f"Permission denied while deleting DB folder: {str(e)}")
        raise HTTPException(status_code=403, detail="Permission denied while deleting DB folder.")
    except Exception as e:
        logger.error(Messages.ERROR_GENERAL.format(error=str(e)))
        raise HTTPException(status_code=500, detail=str(e))

