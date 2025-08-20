# import os
# import uuid
# import chromadb
# from chromadb.utils import embedding_functions
# from utils.enums import ChromaDBConfig, Messages, MetadataKeys
# from utils.logger import setup_logger
# from transformers import AutoTokenizer
# from utils.chunking import chunk_by_headlines # This is the import for your file
# from PyPDF2 import PdfReader
# from docx import Document as DocxDocument
# import tempfile

# logger = setup_logger(__name__)

# # --- Initialize ChromaDB client ---
# persist_dir = os.path.join(os.getcwd(), ChromaDBConfig.DB_DIRECTORY)
# try:
#     if not os.path.isdir(os.path.dirname(persist_dir)):
#         logger.error(f"Invalid database directory path: {persist_dir}")
#         raise RuntimeError(f"Invalid database directory path: {persist_dir}")
#     _client = chromadb.PersistentClient(path=persist_dir)
#     logger.info(f"Initialized ChromaDB client at {persist_dir}")
# except Exception as e:
#     logger.error(f"Failed to initialize ChromaDB client: {str(e)}")
#     raise RuntimeError(f"Failed to initialize ChromaDB client: {str(e)}")

# # --- Initialize embedding function ---
# try:
#     _embedding_func = embedding_functions.SentenceTransformerEmbeddingFunction(
#         model_name=ChromaDBConfig.EMBEDDING_MODEL
#     )
#     logger.info(f"Initialized embedding function: {ChromaDBConfig.EMBEDDING_MODEL}")
# except Exception as e:
#     logger.error(f"Failed to initialize embedding function: {str(e)}")
#     raise RuntimeError(f"Failed to initialize embedding function: {str(e)}")

# # Direct client and collection reference
# client = chromadb.PersistentClient(path=ChromaDBConfig.DB_DIRECTORY)
# embedding_func = embedding_functions.SentenceTransformerEmbeddingFunction(
#     model_name=ChromaDBConfig.EMBEDDING_MODEL
# )
# collection = client.get_or_create_collection(
#     name=ChromaDBConfig.COLLECTION_NAME,
#     embedding_function=embedding_func,
#     metadata={"hnsw:space": "cosine"}
# )

# # Tokenizer for chunking text
# tokenizer = AutoTokenizer.from_pretrained(
#     f"sentence-transformers/{ChromaDBConfig.EMBEDDING_MODEL}"
# )


# def chunk_text_by_tokens(text: str, max_tokens: int = 300, overlap_tokens: int = 50) -> list[str]:
#     """
#     Split a given text into chunks based on token count.

#     This ensures the chunks are suitable for embedding models with token limits.

#     Args:
#         text (str): The full text string to be chunked.
#         max_tokens (int): Maximum number of tokens per chunk.
#         overlap_tokens (int): Number of overlapping tokens between consecutive chunks.

#     Returns:
#         list[str]: A list of text chunks that fit within the token constraints.
#     """
#     tokens = tokenizer.encode(text, add_special_tokens=False)
#     chunks = []
#     start = 0

#     while start < len(tokens):
#         end = start + max_tokens
#         chunk_tokens = tokens[start:end]
#         chunk_text = tokenizer.decode(chunk_tokens, skip_special_tokens=True)
#         chunks.append(chunk_text.strip())
#         start += max_tokens - overlap_tokens

#     return [c for c in chunks if c]


# def get_chroma_collection():
#     """
#     Retrieve or create the ChromaDB collection.

#     Returns:
#         chromadb.api.models.Collection.Collection: The ChromaDB collection object.

#     Raises:
#         RuntimeError: If the collection cannot be retrieved or created.
#     """
#     try:
#         return _client.get_or_create_collection(
#             name=ChromaDBConfig.COLLECTION_NAME,
#             embedding_function=_embedding_func
#         )
#     except Exception as e:
#         logger.error(f"Failed to get or create collection: {str(e)}")
#         raise RuntimeError(f"Failed to get or create collection: {str(e)}")


# def ensure_text_file_exists(file_path: str) -> None:
#     """
#     Ensure the specified text file exists and is not empty.

#     Args:
#         file_path (str): Path to the text file.

#     Raises:
#         FileNotFoundError: If the file does not exist.
#         ValueError: If the file is empty.
#     """
#     if not os.path.exists(file_path):
#         logger.error(Messages.FILE_NOT_FOUND.format(file_path=file_path))
#         raise FileNotFoundError(Messages.FILE_NOT_FOUND.format(file_path=file_path))
#     if os.path.getsize(file_path) == 0:
#         logger.error(f"File '{file_path}' is empty")
#         raise ValueError(f"File '{file_path}' is empty")


# def read_and_chunk_file(file_path: str) -> list[dict]:
#     """
#     Load a text/pdf/docx file and chunk it into structured sections.
    
#     Returns:
#         list[dict]: Each dict contains {"id", "text", "metadata"}.
#     """
#     # Read the text content first
#     extracted_text = ""
#     file_extension = os.path.splitext(file_path)[1].lower()

#     if file_extension == ".txt":
#         with open(file_path, "r", encoding="utf-8") as f:
#             extracted_text = f.read()
#     elif file_extension == ".pdf":
#         reader = PdfReader(file_path)
#         extracted_text = "\n".join([page.extract_text() or "" for page in reader.pages])
#     elif file_extension == ".docx":
#         doc = DocxDocument(file_path)
#         extracted_text = "\n".join([p.text for p in doc.paragraphs])
#     else:
#         logger.error(f"Unsupported file type: {file_extension}")
#         return []

#     # Then pass the extracted text to the chunking function
#     return chunk_by_headlines(extracted_text, os.path.basename(file_path))


# def store_in_collection(collection, chunks: list[str], file_path: str) -> None:
#     """
#     Store a list of text chunks into a ChromaDB collection.

#     This function is deprecated and now uses the new, more robust `sync_text_file_with_collection` function.

#     Args:
#         collection (chromadb.api.models.Collection.Collection): ChromaDB collection object.
#         chunks (list[str]): List of text chunks to store.
#         file_path (str): Source file path for metadata.
#     """
#     logger.info("The `store_in_collection` function is deprecated. Using `sync_text_file_with_collection` instead.")
#     sync_text_file_with_collection(collection, file_path)


# def get_collection_documents(collection) -> dict[str, tuple[str, str]]:
#     """
#     Retrieve all documents from a ChromaDB collection.

#     Args:
#         collection (chromadb.api.models.Collection.Collection): ChromaDB collection object.

#     Returns:
#         dict[str, tuple[str, str]]: Mapping of "source+chunk_index" to (document text, document ID).
#     """
#     data = collection.get(include=['documents', 'metadatas'])
#     return {
#         f"{meta.get(MetadataKeys.SOURCE.value, 'unknown')}{meta.get(MetadataKeys.LINE_NUMBER.value, 0)}": (doc, id)
#         for doc, id, meta in zip(data['documents'], data['ids'], data['metadatas'])
#     }


# def delete_documents(collection, ids: list[str]) -> None:
#     """
#     Delete documents from the collection by their IDs.

#     Args:
#         collection (chromadb.api.models.Collection.Collection): ChromaDB collection object.
#         ids (list[str]): List of document IDs to delete.
#     """
#     if ids:
#         logger.info(f"Deleting {len(ids)} outdated documents")
#         collection.delete(ids=ids)


# def sync_text_file_with_collection(collection, file_path: str, force: bool = False) -> None:
#     """
#     Store a file's content into the ChromaDB collection.
#     Uses section-based chunking instead of token-based.
#     """
#     logger.info(f"Processing file '{file_path}' into collection '{ChromaDBConfig.COLLECTION_NAME}'")

#     ensure_text_file_exists(file_path)
#     # Use the structured chunking from the existing function
#     new_chunks = read_and_chunk_file(file_path)
#     if not new_chunks:
#         logger.info("No new documents to process")
#         return

#     # If force=True -> remove existing chunks for this file
#     if force:
#         logger.info(f"Force mode ON -> deleting old chunks for file: {file_path}")
#         try:
#             collection.delete(where={MetadataKeys.FILENAME.value: os.path.basename(file_path)})
#             logger.info(f"Old chunks for file '{file_path}' deleted successfully")
#         except Exception as e:
#             logger.warning(f"No existing chunks found for file '{file_path}': {str(e)}")

#     # Add new chunks
#     logger.info(Messages.ADDING_DOCUMENTS)
#     collection.add(
#         documents=[c["text"] for c in new_chunks],
#         ids=[c["id"] for c in new_chunks],
#         metadatas=[c["metadata"] for c in new_chunks]
#     )
#     logger.info(Messages.DOCUMENTS_ADDED)


# def initialize_collection():
#     """
#     Initialize the ChromaDB collection and load a default file if empty.

#     If the collection already contains documents, it logs the count.
#     If empty, it loads data from the configured file path.

#     Raises:
#         RuntimeError: If initialization fails.
#     """
#     try:
#         collection = get_chroma_collection()
#         if collection.count() > 0:
#             logger.info(Messages.COLLECTION_EXISTS_WITH_DATA.format(
#                 collection_name=ChromaDBConfig.COLLECTION_NAME,
#                 count=collection.count()
#             ))
#         else:
#             logger.info(Messages.COLLECTION_EXISTS_EMPTY.format(
#                 collection_name=ChromaDBConfig.COLLECTION_NAME
#             ))
#             sync_text_file_with_collection(collection, ChromaDBConfig.FILE_PATH)
#     except Exception as e:
#         logger.error(f"Failed to initialize collection: {str(e)}")
#         raise

# import os
# import uuid
# import chromadb
# from chromadb.utils import embedding_functions
# from utils.enums import ChromaDBConfig, Messages, MetadataKeys
# from utils.logger import setup_logger
# from transformers import AutoTokenizer
# from PyPDF2 import PdfReader
# from docx import Document as DocxDocument
# import tempfile
# import re
# from typing import List, Dict

# logger = setup_logger(__name__)

# # --- Initialize ChromaDB client ---
# persist_dir = os.path.join(os.getcwd(), ChromaDBConfig.DB_DIRECTORY)
# try:
#     if not os.path.isdir(os.path.dirname(persist_dir)):
#         logger.error(f"Invalid database directory path: {persist_dir}")
#         raise RuntimeError(f"Invalid database directory path: {persist_dir}")
#     _client = chromadb.PersistentClient(path=persist_dir)
#     logger.info(f"Initialized ChromaDB client at {persist_dir}")
# except Exception as e:
#     logger.error(f"Failed to initialize ChromaDB client: {str(e)}")
#     raise RuntimeError(f"Failed to initialize ChromaDB client: {str(e)}")

# # --- Initialize embedding function ---
# try:
#     _embedding_func = embedding_functions.SentenceTransformerEmbeddingFunction(
#         model_name=ChromaDBConfig.EMBEDDING_MODEL
#     )
#     logger.info(f"Initialized embedding function: {ChromaDBConfig.EMBEDDING_MODEL}")
# except Exception as e:
#     logger.error(f"Failed to initialize embedding function: {str(e)}")
#     raise RuntimeError(f"Failed to initialize embedding function: {str(e)}")

# # --- Helper function for chunking ---
# def chunk_by_headlines(text: str, filename: str) -> List[Dict]:
#     """
#     Splits text into chunks by headlines and returns them in the required format.
#     The headline is now included in the chunk text to provide more context.
#     """
#     lines = text.splitlines()
#     chunks = []
#     current_headline = None
#     buffer = []
    
#     # This regex is a more robust way to find headlines
#     headline_pattern = re.compile(r"^\s*([A-Z].*)\s*$", re.MULTILINE)

#     for line in lines:
#         line = line.strip()
#         if not line:
#             continue

#         if headline_pattern.match(line) and len(line.split()) < 10:
#             if current_headline and buffer:
#                 # Combine headline and buffered content
#                 chunk_text = f"{current_headline}: {' '.join(buffer).strip()}"
#                 if chunk_text:
#                     chunks.append({
#                         "id": str(uuid.uuid4()),
#                         "text": chunk_text,
#                         "metadata": {
#                             "filename": filename,
#                             "section": current_headline
#                         }
#                     })
#             current_headline = line
#             buffer = []
#         else:
#             buffer.append(line)

#     # Save last chunk
#     if current_headline and buffer:
#         chunk_text = f"{current_headline}: {' '.join(buffer).strip()}"
#         if chunk_text:
#             chunks.append({
#                 "id": str(uuid.uuid4()),
#                 "text": chunk_text,
#                 "metadata": {
#                     "filename": filename,
#                     "section": current_headline
#                 }
#             })

#     # If headline-based chunking yields no results, fall back to simple paragraph splitting
#     if not chunks:
#         paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
#         for i, paragraph in enumerate(paragraphs):
#             chunks.append({
#                 "id": str(uuid.uuid4()),
#                 "text": paragraph,
#                 "metadata": {
#                     "filename": filename,
#                     "section": f"Paragraph {i+1}"
#                 }
#             })

#     return chunks

# # --- Functions for collection management ---
# def get_chroma_collection():
#     """
#     Retrieve or create the ChromaDB collection.
#     """
#     try:
#         return _client.get_or_create_collection(
#             name=ChromaDBConfig.COLLECTION_NAME,
#             embedding_function=_embedding_func
#         )
#     except Exception as e:
#         logger.error(f"Failed to get or create collection: {str(e)}")
#         raise RuntimeError(f"Failed to get or create collection: {str(e)}")

# def ensure_text_file_exists(file_path: str) -> None:
#     """
#     Ensure the specified text file exists and is not empty.
#     """
#     if not os.path.exists(file_path):
#         logger.error(Messages.FILE_NOT_FOUND.format(file_path=file_path))
#         raise FileNotFoundError(Messages.FILE_NOT_FOUND.format(file_path=file_path))
#     if os.path.getsize(file_path) == 0:
#         logger.error(f"File '{file_path}' is empty")
#         raise ValueError(f"File '{file_path}' is empty")


# def read_and_chunk_file(file_path: str) -> list[dict]:
#     """
#     Load a text/pdf/docx file and chunk it into structured sections.
    
#     Returns:
#         list[dict]: Each dict contains {"id", "text", "metadata"}.
#     """
#     # Read the text content first
#     extracted_text = ""
#     file_extension = os.path.splitext(file_path)[1].lower()

#     if file_extension == ".txt":
#         with open(file_path, "r", encoding="utf-8") as f:
#             extracted_text = f.read()
#     elif file_extension == ".pdf":
#         reader = PdfReader(file_path)
#         extracted_text = "\n".join([page.extract_text() or "" for page in reader.pages])
#     elif file_extension == ".docx":
#         doc = DocxDocument(file_path)
#         extracted_text = "\n".join([p.text for p in doc.paragraphs])
#     else:
#         logger.error(f"Unsupported file type: {file_extension}")
#         return []

#     # Then pass the extracted text to the chunking function
#     return chunk_by_headlines(extracted_text, os.path.basename(file_path))


# def sync_text_file_with_collection(collection, file_path: str, force: bool = False) -> None:
#     """
#     Store a file's content into the ChromaDB collection.
#     """
#     logger.info(f"Processing file '{file_path}' into collection '{ChromaDBConfig.COLLECTION_NAME}'")

#     ensure_text_file_exists(file_path)
#     # Use the structured chunking from the existing function
#     new_chunks = read_and_chunk_file(file_path)
#     if not new_chunks:
#         logger.info("No new documents to process")
#         return

#     # If force=True -> remove existing chunks for this file
#     if force:
#         logger.info(f"Force mode ON -> deleting old chunks for file: {file_path}")
#         try:
#             collection.delete(where={MetadataKeys.FILENAME.value: os.path.basename(file_path)})
#             logger.info(f"Old chunks for file '{file_path}' deleted successfully")
#         except Exception as e:
#             logger.warning(f"No existing chunks found for file '{file_path}': {str(e)}")

#     # Add new chunks
#     logger.info(Messages.ADDING_DOCUMENTS)
#     collection.add(
#         documents=[c["text"] for c in new_chunks],
#         ids=[c["id"] for c in new_chunks],
#         metadatas=[c["metadata"] for c in new_chunks]
#     )
#     logger.info(Messages.DOCUMENTS_ADDED)


# def initialize_collection():
#     """
#     Initialize the ChromaDB collection and load a default file if empty.
#     """
#     try:
#         collection = get_chroma_collection()
#         if collection.count() > 0:
#             logger.info(Messages.COLLECTION_EXISTS_WITH_DATA.format(
#                 collection_name=ChromaDBConfig.COLLECTION_NAME,
#                 count=collection.count()
#             ))
#         else:
#             logger.info(Messages.COLLECTION_EXISTS_EMPTY.format(
#                 collection_name=ChromaDBConfig.COLLECTION_NAME
#             ))
#             sync_text_file_with_collection(collection, ChromaDBConfig.FILE_PATH)
#     except Exception as e:
#         logger.error(f"Failed to initialize collection: {str(e)}")
#         raise
import os
import uuid
import chromadb
from chromadb.utils import embedding_functions
from utils.enums import ChromaDBConfig, Messages, MetadataKeys
from utils.logger import setup_logger
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import tempfile
import re
from typing import List, Dict

logger = setup_logger(__name__)

# --- Initialize ChromaDB client ---
persist_dir = os.path.join(os.getcwd(), ChromaDBConfig.DB_DIRECTORY)
try:
    if not os.path.isdir(os.path.dirname(persist_dir)):
        logger.error(f"Invalid database directory path: {persist_dir}")
        raise RuntimeError(f"Invalid database directory path: {persist_dir}")
    _client = chromadb.PersistentClient(path=persist_dir)
    logger.info(f"Initialized ChromaDB client at {persist_dir}")
except Exception as e:
    logger.error(f"Failed to initialize ChromaDB client: {str(e)}")
    raise RuntimeError(f"Failed to initialize ChromaDB client: {str(e)}")

# --- Initialize embedding function ---
try:
    _embedding_func = embedding_functions.SentenceTransformerEmbeddingFunction(
        model_name=ChromaDBConfig.EMBEDDING_MODEL
    )
    logger.info(f"Initialized embedding function: {ChromaDBConfig.EMBEDDING_MODEL}")
except Exception as e:
    logger.error(f"Failed to initialize embedding function: {str(e)}")
    raise RuntimeError(f"Failed to initialize embedding function: {str(e)}")

# --- Helper function for chunking ---
def chunk_by_headlines(text: str, filename: str) -> List[Dict]:
    """
    Splits text into chunks by headlines and returns them in the required format.
    The headline is now included in the chunk text to provide more context.
    """
    lines = text.splitlines()
    chunks = []
    current_headline = None
    buffer = []
    
    # This regex is a more robust way to find headlines
    headline_pattern = re.compile(r"^\s*([A-Z].*)\s*$", re.MULTILINE)

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if headline_pattern.match(line) and len(line.split()) < 10:
            if current_headline and buffer:
                # IMPORTANT: Combine headline and buffered content
                chunk_text = f"{current_headline}: {' '.join(buffer).strip()}"
                if chunk_text:
                    chunks.append({
                        "id": str(uuid.uuid4()),
                        "text": chunk_text,
                        "metadata": {
                            "filename": filename,
                            "section": current_headline
                        }
                    })
            current_headline = line
            buffer = []
        else:
            buffer.append(line)

    # Save last chunk
    if current_headline and buffer:
        # IMPORTANT: Combine headline and buffered content
        chunk_text = f"{current_headline}: {' '.join(buffer).strip()}"
        if chunk_text:
            chunks.append({
                "id": str(uuid.uuid4()),
                "text": chunk_text,
                "metadata": {
                    "filename": filename,
                    "section": current_headline
                }
            })

    # If headline-based chunking yields no results, fall back to simple paragraph splitting
    if not chunks:
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        for i, paragraph in enumerate(paragraphs):
            chunks.append({
                "id": str(uuid.uuid4()),
                "text": paragraph,
                "metadata": {
                    "filename": filename,
                    "section": f"Paragraph {i+1}"
                }
            })

    return chunks

# --- Functions for collection management ---
def get_chroma_collection():
    """
    Retrieve or create the ChromaDB collection.
    """
    try:
        return _client.get_or_create_collection(
            name=ChromaDBConfig.COLLECTION_NAME,
            embedding_function=_embedding_func
        )
    except Exception as e:
        logger.error(f"Failed to get or create collection: {str(e)}")
        raise RuntimeError(f"Failed to get or create collection: {str(e)}")

def ensure_text_file_exists(file_path: str) -> None:
    """
    Ensure the specified text file exists and is not empty.
    """
    if not os.path.exists(file_path):
        logger.error(Messages.FILE_NOT_FOUND.format(file_path=file_path))
        raise FileNotFoundError(Messages.FILE_NOT_FOUND.format(file_path=file_path))
    if os.path.getsize(file_path) == 0:
        logger.error(f"File '{file_path}' is empty")
        raise ValueError(f"File '{file_path}' is empty")


def read_and_chunk_file(file_path: str) -> list[dict]:
    """
    Load a text/pdf/docx file and chunk it into structured sections.
    This version includes a more robust text extraction for .docx files.
    
    Returns:
        list[dict]: Each dict contains {"id", "text", "metadata"}.
    """
    # Read the text content first
    extracted_text = ""
    file_extension = os.path.splitext(file_path)[1].lower()

    if file_extension == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            extracted_text = f.read()
    elif file_extension == ".pdf":
        reader = PdfReader(file_path)
        extracted_text = "\n".join([page.extract_text() or "" for page in reader.pages])
    elif file_extension == ".docx":
        doc = DocxDocument(file_path)
        full_text = []
        for p in doc.paragraphs:
            full_text.append(p.text)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        extracted_text = "\n".join(full_text)
    else:
        logger.error(f"Unsupported file type: {file_extension}")
        return []

    # Then pass the extracted text to the chunking function
    return chunk_by_headlines(extracted_text, os.path.basename(file_path))


def sync_text_file_with_collection(collection, file_path: str, force: bool = False) -> None:
    """
    Store a file's content into the ChromaDB collection.
    """
    logger.info(f"Processing file '{file_path}' into collection '{ChromaDBConfig.COLLECTION_NAME}'")

    ensure_text_file_exists(file_path)
    # Use the structured chunking from the existing function
    new_chunks = read_and_chunk_file(file_path)
    if not new_chunks:
        logger.info("No new documents to process")
        return

    # If force=True -> remove existing chunks for this file
    if force:
        logger.info(f"Force mode ON -> deleting old chunks for file: {file_path}")
        try:
            collection.delete(where={MetadataKeys.FILENAME.value: os.path.basename(file_path)})
            logger.info(f"Old chunks for file '{file_path}' deleted successfully")
        except Exception as e:
            logger.warning(f"No existing chunks found for file '{file_path}': {str(e)}")

    # Add new chunks
    logger.info(Messages.ADDING_DOCUMENTS)
    collection.add(
        documents=[c["text"] for c in new_chunks],
        ids=[c["id"] for c in new_chunks],
        metadatas=[c["metadata"] for c in new_chunks]
    )
    logger.info(Messages.DOCUMENTS_ADDED)


def initialize_collection():
    """
    Initialize the ChromaDB collection and load a default file if empty.
    """
    try:
        collection = get_chroma_collection()
        if collection.count() > 0:
            logger.info(Messages.COLLECTION_EXISTS_WITH_DATA.format(
                collection_name=ChromaDBConfig.COLLECTION_NAME,
                count=collection.count()
            ))
        else:
            logger.info(Messages.COLLECTION_EXISTS_EMPTY.format(
                collection_name=ChromaDBConfig.COLLECTION_NAME
            ))
            sync_text_file_with_collection(collection, ChromaDBConfig.FILE_PATH)
    except Exception as e:
        logger.error(f"Failed to initialize collection: {str(e)}")
        raise
